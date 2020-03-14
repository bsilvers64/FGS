#!/usr/bin/env pyhton 3
from bs4 import BeautifulSoup
import requests
import docx
# added header for fake user-agent
# replace this with user-agent of your system. find it at - https://www.whatismybrowser.com/detect/what-is-my-user-agent
headers = { #enter user-agent here }

url = "https://www.metacritic.com/feature/new-free-games-playstation-xbox-pc-switch"
try:
    res = requests.get(url, headers=headers)
except requests.ConnectionError as e:
    print(e)

row2, row3, k = [], [], 0
doc = docx.Document()
doc.add_heading("list of free PC games available right now -- \n", 2)

page = BeautifulSoup(res.text, 'lxml')
rows = page.find('table', class_='linedtable').find_all('tr')
for i in rows:
    title = i.find_all('td', class_='title')
    rest = i.find_all('td', {'rowspan':'2'})
    for j in rest:
        row3.append(j.get_text().replace(' ', '\n').replace('\n', '').replace('\t', '').replace('\r', '\n'))
    while "" in row3:
        row3.remove("")
    for j in title:
        row = []
        name = j.find('a').get_text()
        row.append(name)
        link = j.find('a').get('href')
        row.append(link)
        li = doc.add_paragraph()
        li.add_run(name).bold = True
        li.add_run(" \n" + link + " \n" + row3[k] + "\n")
        k += 1
        row2.append(row)

#print(row2)
#print(row3)
doc.save("free_games.docx")
